import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from django.conf import settings


def gerar_word(ficha):
    """
    Preenche o template limpo com os dados da ficha
    Mantém logo, cores e formatação original
    """
    # Caminho do template
    template_path = os.path.join(settings.BASE_DIR, 'anamnese', 'template_prontuario.docx')
    
    if not os.path.exists(template_path):
        return gerar_word_basico(ficha)
    
    # Carregar template
    doc = Document(template_path)
    
    # ESTRATÉGIA: Localizar parágrafos específicos e substituir/preencher
    
    # Contador de parágrafo para navegação
    para_idx = 0
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        
        # IDENTIFICAÇÃO
        if 'Data da 1° consulta:' in texto:
            # Adicionar data ao final
            para.add_run(f" {ficha.data_consulta.strftime('%d/%m/%Y')}")
        
        elif 'Nome:' in texto and 'Data de nascimento' not in texto:
            para.add_run(f" {ficha.nome}")
        
        elif 'Endereço:' in texto:
            para.add_run(f" {ficha.endereco or ''}")
        
        elif 'Bairro:' in texto:
            # Substituir o texto mantendo a estrutura
            para.clear()
            run = para.add_run(f"Bairro: {ficha.bairro or ''}                                               Email: {ficha.email}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Profissão:' in texto and 'Celular:' in texto:
            para.clear()
            run = para.add_run(f"Profissão: {ficha.profissao}                                          Celular: {ficha.celular}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Data de nascimento:' in texto:
            sexo_m = 'x' if ficha.sexo == 'M' else ' '
            sexo_f = 'x' if ficha.sexo == 'F' else ' '
            para.clear()
            run = para.add_run(f"Data de nascimento: {ficha.data_nascimento.strftime('%d/%m/%Y')}    Idade: {ficha.idade}     Sexo: ({sexo_m}) Masculino  ({sexo_f}) Feminino")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif texto == 'Motivo da consulta:':
            para.add_run(f" {ficha.motivo_consulta}")
        
        elif texto == 'Observações:':
            para.add_run(f" {ficha.observacoes_iniciais or ''}")
        
        # HISTÓRICO SOCIAL
        elif 'Estado civil:' in texto and 'Composição familiar:' in texto:
            para.clear()
            run = para.add_run(f"Estado civil: {ficha.get_estado_civil_display()}                             Composição familiar: {ficha.composicao_familiar or ''}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Quem compra os alimentos:' in texto:
            para.clear()
            run = para.add_run(f"Quem compra os alimentos: {ficha.quem_compra_alimentos or ''}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'A compra e feita:' in texto:
            diario = 'x' if ficha.frequencia_compra == 'diariamente' else ' '
            semanal = 'x' if ficha.frequencia_compra == 'semanalmente' else ' '
            mensal = 'x' if ficha.frequencia_compra == 'mensalmente' else ' '
            para.clear()
            run = para.add_run(f"A compra e feita: ({diario}) diariamente ({semanal}) semanalmente ({mensal}) mensalmente")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Quem prepara as refeições:' in texto:
            para.clear()
            run = para.add_run(f"Quem prepara as refeições: {ficha.quem_prepara_refeicoes or ''}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Com quem realiza as refeições:' in texto:
            para.clear()
            run = para.add_run(f"Com quem realiza as refeições: {ficha.com_quem_faz_refeicoes or ''}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Faz uso de bebidas alcoólicas?' in texto:
            para.clear()
            run = para.add_run(f"Faz uso de bebidas alcoólicas? Frequências: {ficha.bebidas_alcoolicas or 'Não'}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Fuma ou já fumou?' in texto:
            para.clear()
            run = para.add_run(f"Fuma ou já fumou? {ficha.fuma or 'Não'}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        # OBJETIVO
        elif 'Por que procurou atendimento:' in texto:
            para.clear()
            run = para.add_run(f"Por que procurou atendimento: {ficha.porque_procurou}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
        
        elif 'Meta pessoal:' in texto:
            para.clear()
            run = para.add_run(f"Meta pessoal: {ficha.meta_pessoal}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
    
    # PREENCHER TABELA 1: DADOS CLÍNICOS (Sintomas)
    if len(doc.tables) >= 1:
        tabela_sintomas = doc.tables[0]
        
        sintomas = [
            (2, ficha.vomito, ficha.vomito_obs or ''),
            (3, ficha.nausea, ficha.nausea_obs or ''),
            (4, ficha.mastigacao, ficha.mastigacao_obs or ''),
            (5, ficha.degluticao, ficha.degluticao_obs or ''),
            (6, ficha.digestao, ficha.digestao_obs or ''),
            (7, ficha.pirose, ficha.pirose_obs or ''),
            (8, ficha.refluxo, ficha.refluxo_obs or ''),
            (9, ficha.diarreia, ficha.diarreia_obs or ''),
            (10, ficha.obstipacao, ficha.obstipacao_obs or ''),
            (11, ficha.insonia, ficha.insonia_obs or ''),
            (12, ficha.estresse, ficha.estresse_obs or ''),
            (13, ficha.cansaco, ficha.cansaco_obs or ''),
            (14, ficha.ansiedade, ficha.ansiedade_obs or ''),
            (15, ficha.depressao, ficha.depressao_obs or ''),
            (16, ficha.trabalho_estudo_afeta, ficha.trabalho_estudo_obs or ''),
        ]
        
        for row_idx, tem_sintoma, obs in sintomas:
            if row_idx < len(tabela_sintomas.rows):
                row = tabela_sintomas.rows[row_idx]
                if len(row.cells) >= 4:
                    # Limpar e preencher
                    row.cells[1].text = 'Sim' if tem_sintoma else ''
                    row.cells[2].text = '' if tem_sintoma else 'Não'
                    row.cells[3].text = obs
    
    # PREENCHER TABELA 2: Outros dados clínicos
    if len(doc.tables) >= 2:
        tabela_outros = doc.tables[1]
        
        # Linha 0: Lesões de pele
        if len(tabela_outros.rows) > 0:
            tabela_outros.rows[0].cells[0].text = f"Possui lesões, problemas na pele cabelo ou unha? {ficha.lesoes_pele or ''}"
        
        # Linha 1: Cirurgia
        if len(tabela_outros.rows) > 1:
            tabela_outros.rows[1].cells[0].text = f"Já passou por algum tipo de cirurgia? {ficha.cirurgia or 'Não               Qual?                            Quando?'}"
        
        # Linha 2: Hábito intestinal
        if len(tabela_outros.rows) > 2:
            diario = 'x' if ficha.habito_intestinal == 'diario' else ' '
            ate3 = 'x' if ficha.habito_intestinal == 'ate_3_dias' else ' '
            mais3 = 'x' if ficha.habito_intestinal == 'mais_3_dias' else ' '
            tabela_outros.rows[2].cells[0].text = f"Hábito intestinal: ({diario}) Diário ({ate3}) Até 3 dias  ({mais3}) Mais de 3 dias  (  ) Outros"
        
        # Linha 3: Consistência
        if len(tabela_outros.rows) > 3:
            normal = 'x' if ficha.consistencia_fezes == 'normal' else ' '
            amolec = 'x' if ficha.consistencia_fezes == 'amolecidas' else ' '
            duras = 'x' if ficha.consistencia_fezes == 'duras' else ' '
            tabela_outros.rows[3].cells[0].text = f"Consistência das fezes: ({normal}) Normal ({amolec}) Amolecidas ({duras}) Duras"
        
        # Linha 4: Diurese
        if len(tabela_outros.rows) > 4:
            tabela_outros.rows[4].cells[0].text = f"Diurese (Quantidade/Coloração): {ficha.diurese or ''}"
        
        # Linha 5: Patologia
        if len(tabela_outros.rows) > 5:
            tabela_outros.rows[5].cells[0].text = f"Possui alguma patologia? {ficha.patologia or 'Não                         Qual?                             Desde quando?'}"
        
        # Linha 6: Medicamento
        if len(tabela_outros.rows) > 6:
            tabela_outros.rows[6].cells[0].text = f"Toma algum tipo de medicamento? {ficha.medicamento or 'Não                       Qual?                            Tempo?'}"
    
    # PREENCHER TABELA 3: ANTROPOMETRIA
    if len(doc.tables) >= 3:
        tabela_antro = doc.tables[2]
        
        # Preencher dados básicos (linhas 2-4)
        if len(tabela_antro.rows) > 2:
            tabela_antro.rows[2].cells[0].text = ficha.data_consulta.strftime('%d/%m/%Y')
            tabela_antro.rows[2].cells[3].text = str(ficha.peso)
        
        if len(tabela_antro.rows) > 3:
            tabela_antro.rows[3].cells[0].text = ficha.data_consulta.strftime('%d/%m/%Y')
            tabela_antro.rows[3].cells[3].text = str(ficha.estatura)
        
        if len(tabela_antro.rows) > 4:
            tabela_antro.rows[4].cells[0].text = ficha.data_consulta.strftime('%d/%m/%Y')
            tabela_antro.rows[4].cells[3].text = str(ficha.imc) if ficha.imc else ''
        
        # Preencher circunferências (linhas 5-16)
        circunferencias = [
            (5, ficha.circunferencia_braco_esq),
            (6, ficha.circunferencia_braco_dir),
            (7, ficha.circunferencia_peitoral),
            (8, ficha.circunferencia_cintura),
            (9, ficha.circunferencia_abdominal),
            (10, ficha.circunferencia_quadril),
            (11, ficha.circunferencia_coxa_dir),
            (12, ficha.circunferencia_coxa_esq),
            (13, ficha.circunferencia_panturrilha_esq),
            (14, ficha.circunferencia_panturrilha_dir),
            (15, ficha.circunferencia_pescoco),
            (16, ficha.relacao_cq),
        ]
        
        for row_idx, valor in circunferencias:
            if row_idx < len(tabela_antro.rows) and valor:
                tabela_antro.rows[row_idx].cells[0].text = ficha.data_consulta.strftime('%d/%m/%Y')
                tabela_antro.rows[row_idx].cells[3].text = str(valor)
        
        # Preencher dobras cutâneas (linhas 17-25)
        dobras = [
            (17, ficha.dc_triceps),
            (18, ficha.dc_biceps),
            (19, ficha.dc_escapula),
            (20, ficha.dc_suprailiaca),
            (21, ficha.dc_axilar_media),
            (22, ficha.dc_peitoral),
            (23, ficha.dc_abdominal),
            (24, ficha.dc_coxa_media),
            (25, ficha.dc_panturrilha),
        ]
        
        for row_idx, valor in dobras:
            if row_idx < len(tabela_antro.rows) and valor:
                tabela_antro.rows[row_idx].cells[3].text = str(valor)
        
        # % Gordura (linha 26)
        if len(tabela_antro.rows) > 26 and ficha.percentual_gordura:
            tabela_antro.rows[26].cells[3].text = str(ficha.percentual_gordura)
    
    # Adicionar diagnóstico e orientações no final
    if ficha.diagnostico_problema or ficha.orientacoes_texto:
        # Adicionar quebra de página
        doc.add_page_break()
        
        if ficha.diagnostico_problema:
            p = doc.add_paragraph()
            run = p.add_run('Diagnóstico nutricional')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            
            texto_diag = f"Problema: {ficha.diagnostico_problema}\n"
            if ficha.diagnostico_etiologia:
                texto_diag += f"Etiologia: {ficha.diagnostico_etiologia}\n"
            if ficha.diagnostico_sinais:
                texto_diag += f"Sinais/Sintomas: {ficha.diagnostico_sinais}"
            
            doc.add_paragraph(texto_diag)
        
        if ficha.orientacoes_texto:
            p = doc.add_paragraph()
            run = p.add_run('Orientações e devolutivas')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            
            if ficha.data_orientacao:
                doc.add_paragraph(f"Data da orientação: {ficha.data_orientacao.strftime('%d/%m/%Y')}")
            
            doc.add_paragraph(ficha.orientacoes_texto)
            
            if ficha.calculos_observacoes:
                doc.add_paragraph(ficha.calculos_observacoes)
    
    # Salvar documento
    output_dir = os.path.join(settings.MEDIA_ROOT, 'docs')
    os.makedirs(output_dir, exist_ok=True)
    filename = f'prontuario_{ficha.id}_{ficha.nome.replace(" ", "_")}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filepath


def gerar_word_basico(ficha):
    """Fallback caso o template não exista"""
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    
    title = doc.add_paragraph()
    title_run = title.add_run('FICHA DE ANAMNESE ALIMENTAR')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Nome: {ficha.nome}")
    doc.add_paragraph(f"Email: {ficha.email}")
    doc.add_paragraph(f"Peso: {ficha.peso} kg")
    doc.add_paragraph(f"Altura: {ficha.estatura} m")
    doc.add_paragraph(f"IMC: {ficha.imc}")
    
    output_dir = os.path.join(settings.MEDIA_ROOT, 'docs')
    os.makedirs(output_dir, exist_ok=True)
    filename = f'prontuario_{ficha.id}_{ficha.nome.replace(" ", "_")}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filepath


def gerar_pdf(ficha):
    """Gera PDF simples"""
    pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
    os.makedirs(pdf_dir, exist_ok=True)
    
    filename = f'prontuario_{ficha.id}_{ficha.nome.replace(" ", "_")}.pdf'
    filepath = os.path.join(pdf_dir, filename)
    
    doc = SimpleDocTemplate(filepath, pagesize=A4,
                           topMargin=2*cm, bottomMargin=2*cm,
                           leftMargin=2*cm, rightMargin=2*cm)
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16,
                                textColor=colors.HexColor('#2c3e50'), spaceAfter=20,
                                alignment=TA_CENTER, fontName='Helvetica-Bold')
    
    normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=10,
                                  leading=14, spaceAfter=8)
    
    story = []
    
    story.append(Paragraph("FICHA DE ANAMNESE ALIMENTAR", title_style))
    story.append(Spacer(1, 0.5*cm))
    
    story.append(Paragraph(f"<b>Nome:</b> {ficha.nome}", normal_style))
    story.append(Paragraph(f"<b>Email:</b> {ficha.email}", normal_style))
    story.append(Paragraph(f"<b>Peso:</b> {ficha.peso} kg", normal_style))
    story.append(Paragraph(f"<b>Altura:</b> {ficha.estatura} m", normal_style))
    story.append(Paragraph(f"<b>IMC:</b> {ficha.imc}", normal_style))
    
    story.append(Spacer(1, 1*cm))
    data_geracao = f"Documento gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, 
                                  textColor=colors.grey, alignment=TA_CENTER)
    story.append(Paragraph(data_geracao, footer_style))
    
    doc.build(story)
    return filepath